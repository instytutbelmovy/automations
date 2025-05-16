from pathlib import Path
from docx import Document
from docx.document import Document as DocxDocument
from .doc_reader import DocReader, SourceDocument


class DocxReader(DocReader):
    def read(self, file_path: str | Path) -> SourceDocument:
        """Чытае DOCX файл па шляху і вяртае SourceDocument з метададзенымі і параграфамі."""
        doc: DocxDocument = Document(str(file_path))

        # Збіраем метададзеныя
        title = doc.core_properties.title if doc.core_properties.title else "Без назвы"
        author = doc.core_properties.author if doc.core_properties.author else None
        language = doc.core_properties.language if doc.core_properties.language else None
        date = str(doc.core_properties.created) if doc.core_properties.created else None

        document = SourceDocument(title=title, author=author, language=language, publication_date=date)

        # Апрацоўваем змесціва
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Пропускаем пустыя параграфы
                document.paragraphs.append(text)

        return document
